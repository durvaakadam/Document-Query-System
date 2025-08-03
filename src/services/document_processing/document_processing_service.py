import asyncio
import re
import logging
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from typing import Dict, List, Optional, Tuple, Any
from pathlib import Path
import json
from datetime import datetime

# Document processing libraries
import fitz  # pymupdf
import pdfplumber
from docx import Document
import email
from email import policy
from email.mime.text import MIMEText
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords

# OCR libraries (optional, for scanned PDFs)
try:
    import pytesseract
    from PIL import Image
    import pdf2image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logging.warning("OCR libraries not available. Install pytesseract, Pillow, and pdf2image for OCR support.")

# Download NLTK data if not present
try:
    nltk.data.find('tokenizers/punkt')
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('punkt')
    nltk.download('stopwords')

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class TextChunk:
    """Represents a chunk of extracted text with metadata"""
    content: str
    chunk_id: str
    chunk_type: str  # 'section', 'paragraph', 'table', 'list'
    page_number: Optional[int] = None
    section_title: Optional[str] = None
    start_position: Optional[int] = None
    end_position: Optional[int] = None
    metadata: Dict[str, Any] = field(default_factory=dict)
    word_count: int = 0
    
    def __post_init__(self):
        if not self.word_count:
            self.word_count = len(self.content.split())

@dataclass
class ExtractedDocument:
    """Complete extracted document with all processed content"""
    document_id: str
    source_file: str
    file_type: str
    full_text: str
    chunks: List[TextChunk]
    tables: List[Dict] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    extraction_timestamp: datetime = field(default_factory=datetime.now)
    processing_stats: Dict[str, Any] = field(default_factory=dict)

class BaseTextExtractor(ABC):
    """Abstract base class for text extractors"""
    
    @abstractmethod
    async def extract_text(self, file_path: str) -> ExtractedDocument:
        """Extract text from document"""
        pass
    
    @abstractmethod
    def validate_file(self, file_path: str) -> bool:
        """Validate if file can be processed"""
        pass

class PDFTextExtractor(BaseTextExtractor):
    """Advanced PDF text extraction with multiple strategies"""
    
    def __init__(self, enable_ocr: bool = True, ocr_language: str = 'eng'):
        self.enable_ocr = enable_ocr and OCR_AVAILABLE
        self.ocr_language = ocr_language
    
    def validate_file(self, file_path: str) -> bool:
        """Validate PDF file"""
        try:
            with fitz.open(file_path) as doc:
                return doc.page_count > 0
        except Exception as e:
            logger.error(f"PDF validation failed: {e}")
            return False
    
    async def extract_text(self, file_path: str) -> ExtractedDocument:
        """Extract text from PDF using multiple strategies"""
        logger.info(f"Extracting text from PDF: {file_path}")
        
        # Try PyMuPDF first (fastest and most reliable)
        try:
            extracted_doc = await self._extract_with_pymupdf(file_path)
            
            # Check if we got meaningful text
            if self._has_sufficient_text(extracted_doc.full_text):
                logger.info("PyMuPDF extraction successful")
                return extracted_doc
            else:
                logger.warning("PyMuPDF extracted insufficient text, trying pdfplumber")
                
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed: {e}")
        
        # Try pdfplumber for complex layouts
        try:
            extracted_doc = await self._extract_with_pdfplumber(file_path)
            
            if self._has_sufficient_text(extracted_doc.full_text):
                logger.info("pdfplumber extraction successful")
                return extracted_doc
            else:
                logger.warning("pdfplumber extracted insufficient text")
                
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {e}")
        
        # Try OCR as last resort
        if self.enable_ocr:
            try:
                logger.info("Attempting OCR extraction")
                extracted_doc = await self._extract_with_ocr(file_path)
                return extracted_doc
            except Exception as e:
                logger.error(f"OCR extraction failed: {e}")
        
        # If all methods fail, return empty document
        return ExtractedDocument(
            document_id=Path(file_path).stem,
            source_file=file_path,
            file_type="pdf",
            full_text="",
            chunks=[],
            processing_stats={"extraction_method": "failed", "error": "All extraction methods failed"}
        )
    
    def _has_sufficient_text(self, text: str, min_words: int = 50) -> bool:
        """Check if extracted text is sufficient"""
        return len(text.split()) >= min_words
    
    async def _extract_with_pymupdf(self, file_path: str) -> ExtractedDocument:
        """Extract text using PyMuPDF"""
        doc = fitz.open(file_path)
        full_text = ""
        chunks = []
        
        for page_num in range(doc.page_count):
            page = doc[page_num]
            page_text = page.get_text()
            
            if page_text.strip():
                full_text += page_text + "\n"
                
                # Create page-level chunks
                page_chunks = await self._create_page_chunks(
                    page_text, page_num + 1, f"{Path(file_path).stem}_pymupdf"
                )
                chunks.extend(page_chunks)
        
        doc.close()
        
        return ExtractedDocument(
            document_id=Path(file_path).stem,
            source_file=file_path,
            file_type="pdf",
            full_text=full_text,
            chunks=chunks,
            metadata={"total_pages": doc.page_count},
            processing_stats={"extraction_method": "pymupdf", "pages_processed": doc.page_count}
        )
    
    async def _extract_with_pdfplumber(self, file_path: str) -> ExtractedDocument:
        """Extract text using pdfplumber (better for tables and complex layouts)"""
        full_text = ""
        chunks = []
        tables = []
        
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                # Extract text
                page_text = page.extract_text()
                if page_text:
                    full_text += page_text + "\n"
                    
                    # Create page-level chunks
                    page_chunks = await self._create_page_chunks(
                        page_text, page_num + 1, f"{Path(file_path).stem}_pdfplumber"
                    )
                    chunks.extend(page_chunks)
                
                # Extract tables
                page_tables = page.extract_tables()
                for table_idx, table in enumerate(page_tables):
                    if table:
                        table_data = {
                            "page": page_num + 1,
                            "table_index": table_idx,
                            "headers": table[0] if table else [],
                            "rows": table[1:] if len(table) > 1 else [],
                            "raw_data": table
                        }
                        tables.append(table_data)
                        
                        # Create table chunk
                        table_text = self._table_to_text(table)
                        table_chunk = TextChunk(
                            content=table_text,
                            chunk_id=f"{Path(file_path).stem}_table_{page_num}_{table_idx}",
                            chunk_type="table",
                            page_number=page_num + 1,
                            metadata={"table_data": table_data}
                        )
                        chunks.append(table_chunk)
        
        return ExtractedDocument(
            document_id=Path(file_path).stem,
            source_file=file_path,
            file_type="pdf",
            full_text=full_text,
            chunks=chunks,
            tables=tables,
            metadata={"total_pages": len(pdf.pages)},
            processing_stats={"extraction_method": "pdfplumber", "tables_found": len(tables)}
        )
    
    async def _extract_with_ocr(self, file_path: str) -> ExtractedDocument:
        """Extract text using OCR for scanned PDFs"""
        if not OCR_AVAILABLE:
            raise RuntimeError("OCR libraries not available")
        
        # Convert PDF to images
        images = pdf2image.convert_from_path(file_path)
        full_text = ""
        chunks = []
        
        for page_num, image in enumerate(images):
            # Perform OCR
            page_text = pytesseract.image_to_string(image, lang=self.ocr_language)
            
            if page_text.strip():
                full_text += page_text + "\n"
                
                # Create page-level chunks
                page_chunks = await self._create_page_chunks(
                    page_text, page_num + 1, f"{Path(file_path).stem}_ocr"
                )
                chunks.extend(page_chunks)
        
        return ExtractedDocument(
            document_id=Path(file_path).stem,
            source_file=file_path,
            file_type="pdf",
            full_text=full_text,
            chunks=chunks,
            processing_stats={"extraction_method": "ocr", "pages_processed": len(images)}
        )
    
    async def _create_page_chunks(self, page_text: str, page_num: int, doc_id: str) -> List[TextChunk]:
        """Create semantic chunks from page text"""
        chunks = []
        
        # Split by sections if we can identify them
        sections = self._identify_sections(page_text)
        
        if sections:
            for idx, (section_title, section_text) in enumerate(sections):
                chunk = TextChunk(
                    content=section_text,
                    chunk_id=f"{doc_id}_p{page_num}_s{idx}",
                    chunk_type="section",
                    page_number=page_num,
                    section_title=section_title
                )
                chunks.append(chunk)
        else:
            # Fall back to paragraph-based chunking
            paragraphs = [p.strip() for p in page_text.split('\n\n') if p.strip()]
            
            for idx, paragraph in enumerate(paragraphs):
                if len(paragraph) > 50:  # Only meaningful paragraphs
                    chunk = TextChunk(
                        content=paragraph,
                        chunk_id=f"{doc_id}_p{page_num}_para{idx}",
                        chunk_type="paragraph",
                        page_number=page_num
                    )
                    chunks.append(chunk)
        
        return chunks
    
    def _identify_sections(self, text: str) -> List[Tuple[str, str]]:
        """Identify sections in insurance/legal documents"""
        sections = []
        
        # Common patterns for insurance/legal document sections
        section_patterns = [
            r'^([A-Z][A-Z\s]+):',  # ALL CAPS headers like "COVERAGE:"
            r'^(\d+\.\s+[A-Z][^.]+):',  # Numbered sections like "1. DEFINITIONS:"
            r'^([A-Z][a-z]+\s+\d+[.:]\s*[A-Z][^.]+)',  # "Section 1: Coverage Details"
            r'^(\([a-z]\)\s*[A-Z][^.]+)',  # "(a) Premium Payment"
        ]
        
        lines = text.split('\n')
        current_section = None
        current_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if this line is a section header
            is_header = False
            for pattern in section_patterns:
                match = re.match(pattern, line)
                if match:
                    # Save previous section
                    if current_section and current_content:
                        sections.append((current_section, '\n'.join(current_content)))
                    
                    # Start new section
                    current_section = match.group(1)
                    current_content = []
                    is_header = True
                    break
            
            if not is_header and current_section:
                current_content.append(line)
        
        # Add the last section
        if current_section and current_content:
            sections.append((current_section, '\n'.join(current_content)))
        
        return sections
    
    def _table_to_text(self, table: List[List[str]]) -> str:
        """Convert table data to readable text"""
        if not table:
            return ""
        
        text_parts = []
        headers = table[0] if table else []
        
        if headers:
            text_parts.append("Table Headers: " + " | ".join(str(h) for h in headers if h))
        
        for row in table[1:]:
            if row and any(cell for cell in row if cell):
                row_text = " | ".join(str(cell) for cell in row if cell)
                text_parts.append(row_text)
        
        return "\n".join(text_parts)

class DOCXTextExtractor(BaseTextExtractor):
    """DOCX document text extraction"""
    
    def validate_file(self, file_path: str) -> bool:
        """Validate DOCX file"""
        try:
            Document(file_path)
            return True
        except Exception as e:
            logger.error(f"DOCX validation failed: {e}")
            return False
    
    async def extract_text(self, file_path: str) -> ExtractedDocument:
        """Extract text from DOCX file"""
        logger.info(f"Extracting text from DOCX: {file_path}")
        
        doc = Document(file_path)
        full_text = ""
        chunks = []
        tables = []
        
        # Extract paragraphs
        for para_idx, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                full_text += paragraph.text + "\n"
                
                # Determine chunk type based on style
                chunk_type = "paragraph"
                section_title = None
                
                if paragraph.style.name.startswith('Heading'):
                    chunk_type = "section"
                    section_title = paragraph.text.strip()
                
                chunk = TextChunk(
                    content=paragraph.text,
                    chunk_id=f"{Path(file_path).stem}_para_{para_idx}",
                    chunk_type=chunk_type,
                    section_title=section_title,
                    metadata={"style": paragraph.style.name}
                )
                chunks.append(chunk)
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            table_text_parts = []
            
            for row_idx, row in enumerate(table.rows):
                row_data = []
                row_text = []
                
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                    row_text.append(cell_text)
                
                table_data.append(row_data)
                table_text_parts.append(" | ".join(row_text))
            
            if table_data:
                table_info = {
                    "table_index": table_idx,
                    "headers": table_data[0] if table_data else [],
                    "rows": table_data[1:] if len(table_data) > 1 else [],
                    "raw_data": table_data
                }
                tables.append(table_info)
                
                # Create table chunk
                table_text = "\n".join(table_text_parts)
                table_chunk = TextChunk(
                    content=table_text,
                    chunk_id=f"{Path(file_path).stem}_table_{table_idx}",
                    chunk_type="table",
                    metadata={"table_data": table_info}
                )
                chunks.append(table_chunk)
                full_text += table_text + "\n"
        
        return ExtractedDocument(
            document_id=Path(file_path).stem,
            source_file=file_path,
            file_type="docx",
            full_text=full_text,
            chunks=chunks,
            tables=tables,
            processing_stats={"extraction_method": "python-docx", "paragraphs": len(doc.paragraphs), "tables": len(tables)}
        )

class EmailTextExtractor(BaseTextExtractor):
    """Email text extraction"""
    
    def validate_file(self, file_path: str) -> bool:
        """Validate email file"""
        try:
            with open(file_path, 'rb') as f:
                email.message_from_bytes(f.read(), policy=policy.default)
            return True
        except Exception as e:
            logger.error(f"Email validation failed: {e}")
            return False
    
    async def extract_text(self, file_path: str) -> ExtractedDocument:
        """Extract text from email file"""
        logger.info(f"Extracting text from email: {file_path}")
        
        with open(file_path, 'rb') as f:
            msg = email.message_from_bytes(f.read(), policy=policy.default)
        
        # Extract headers
        headers = {
            'subject': msg.get('Subject', ''),
            'from': msg.get('From', ''),
            'to': msg.get('To', ''),
            'date': msg.get('Date', ''),
            'cc': msg.get('Cc', ''),
            'bcc': msg.get('Bcc', '')
        }
        
        # Extract body
        body_text = ""
        chunks = []
        
        if msg.is_multipart():
            for part in msg.walk():
                if part.get_content_type() == "text/plain":
                    try:
                        body_text += part.get_content()
                    except Exception as e:
                        logger.warning(f"Failed to extract email part: {e}")
        else:
            if msg.get_content_type() == "text/plain":
                body_text = msg.get_content()
        
        # Create header chunk
        header_text = f"Subject: {headers['subject']}\nFrom: {headers['from']}\nTo: {headers['to']}\nDate: {headers['date']}"
        header_chunk = TextChunk(
            content=header_text,
            chunk_id=f"{Path(file_path).stem}_headers",
            chunk_type="headers",
            metadata=headers
        )
        chunks.append(header_chunk)
        
        # Create body chunks
        if body_text:
            paragraphs = [p.strip() for p in body_text.split('\n\n') if p.strip()]
            
            for idx, paragraph in enumerate(paragraphs):
                if len(paragraph) > 20:  # Only meaningful paragraphs
                    chunk = TextChunk(
                        content=paragraph,
                        chunk_id=f"{Path(file_path).stem}_body_{idx}",
                        chunk_type="paragraph"
                    )
                    chunks.append(chunk)
        
        full_text = header_text + "\n\n" + body_text
        
        return ExtractedDocument(
            document_id=Path(file_path).stem,
            source_file=file_path,
            file_type="email",
            full_text=full_text,
            chunks=chunks,
            metadata=headers,
            processing_stats={"extraction_method": "email", "is_multipart": msg.is_multipart()}
        )

class SmartTextChunker:
    """Advanced text chunking for insurance/legal documents"""
    
    def __init__(self, 
                 chunk_size: int = 1000,
                 overlap: int = 200,
                 min_chunk_size: int = 100):
        self.chunk_size = chunk_size
        self.overlap = overlap
        self.min_chunk_size = min_chunk_size
        self.stop_words = set(stopwords.words('english'))
    
    async def rechunk_document(self, extracted_doc: ExtractedDocument) -> ExtractedDocument:
        """Apply advanced chunking strategies to improve semantic coherence"""
        logger.info(f"Applying advanced chunking to document: {extracted_doc.document_id}")
        
        # Combine similar chunks and optimize for semantic search
        optimized_chunks = []
        
        # Group chunks by type
        section_chunks = [c for c in extracted_doc.chunks if c.chunk_type == "section"]
        paragraph_chunks = [c for c in extracted_doc.chunks if c.chunk_type == "paragraph"]
        table_chunks = [c for c in extracted_doc.chunks if c.chunk_type == "table"]
        
        # Process sections with intelligent splitting
        for chunk in section_chunks:
            if len(chunk.content) > self.chunk_size:
                sub_chunks = await self._split_large_chunk(chunk)
                optimized_chunks.extend(sub_chunks)
            else:
                optimized_chunks.append(chunk)
        
        # Combine small consecutive paragraphs
        combined_para_chunks = await self._combine_small_chunks(paragraph_chunks)
        optimized_chunks.extend(combined_para_chunks)
        
        # Keep table chunks as-is (they're usually well-structured)
        optimized_chunks.extend(table_chunks)
        
        # Create sliding window chunks for better coverage
        sliding_chunks = await self._create_sliding_window_chunks(extracted_doc.full_text, extracted_doc.document_id)
        optimized_chunks.extend(sliding_chunks)
        
        # Update the document with optimized chunks
        extracted_doc.chunks = optimized_chunks
        extracted_doc.processing_stats["chunking_strategy"] = "advanced"
        extracted_doc.processing_stats["total_chunks"] = len(optimized_chunks)
        
        return extracted_doc
    
    async def _split_large_chunk(self, chunk: TextChunk) -> List[TextChunk]:
        """Split large chunks while preserving semantic boundaries"""
        sentences = sent_tokenize(chunk.content)
        sub_chunks = []
        current_chunk = ""
        chunk_idx = 0
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) > self.chunk_size and current_chunk:
                # Create sub-chunk
                sub_chunk = TextChunk(
                    content=current_chunk.strip(),
                    chunk_id=f"{chunk.chunk_id}_sub_{chunk_idx}",
                    chunk_type=chunk.chunk_type,
                    page_number=chunk.page_number,
                    section_title=chunk.section_title,
                    metadata=chunk.metadata.copy()
                )
                sub_chunks.append(sub_chunk)
                
                # Start new chunk with overlap
                current_chunk = sentence + " "
                chunk_idx += 1
            else:
                current_chunk += sentence + " "
        
        # Add the last chunk
        if current_chunk.strip():
            sub_chunk = TextChunk(
                content=current_chunk.strip(),
                chunk_id=f"{chunk.chunk_id}_sub_{chunk_idx}",
                chunk_type=chunk.chunk_type,
                page_number=chunk.page_number,
                section_title=chunk.section_title,
                metadata=chunk.metadata.copy()
            )
            sub_chunks.append(sub_chunk)
        
        return sub_chunks
    
    async def _combine_small_chunks(self, chunks: List[TextChunk]) -> List[TextChunk]:
        """Combine small consecutive chunks"""
        if not chunks:
            return []
        
        combined_chunks = []
        current_combined = None
        
        for chunk in chunks:
            if len(chunk.content) < self.min_chunk_size:
                if current_combined is None:
                    current_combined = TextChunk(
                        content=chunk.content,
                        chunk_id=f"{chunk.chunk_id}_combined",
                        chunk_type="paragraph",
                        page_number=chunk.page_number,
                        metadata=chunk.metadata.copy()
                    )
                else:
                    current_combined.content += "\n\n" + chunk.content
                    current_combined.chunk_id += f"_{chunk.chunk_id.split('_')[-1]}"
            else:
                # Add the current combined chunk if it exists
                if current_combined:
                    combined_chunks.append(current_combined)
                    current_combined = None
                
                # Add the current chunk
                combined_chunks.append(chunk)
        
        # Add the last combined chunk
        if current_combined:
            combined_chunks.append(current_combined)
        
        return combined_chunks
    
    async def _create_sliding_window_chunks(self, full_text: str, doc_id: str) -> List[TextChunk]:
        """Create sliding window chunks for comprehensive coverage"""
        words = word_tokenize(full_text.lower())
        # Remove stop words for better semantic content
        meaningful_words = [w for w in words if w.isalnum() and w not in self.stop_words]
        
        if len(meaningful_words) < self.chunk_size:
            return []
        
        sliding_chunks = []
        step_size = self.chunk_size - self.overlap
        
        for i in range(0, len(meaningful_words) - self.chunk_size, step_size):
            chunk_words = meaningful_words[i:i + self.chunk_size]
            chunk_content = " ".join(chunk_words)
            
            chunk = TextChunk(
                content=chunk_content,
                chunk_id=f"{doc_id}_sliding_{i // step_size}",
                chunk_type="sliding_window",
                metadata={"start_word": i, "end_word": i + self.chunk_size}
            )
            sliding_chunks.append(chunk)
        
        return sliding_chunks

class DocumentTextProcessor:
    """Main text processing service that orchestrates all extractors"""
    
    def __init__(self, enable_ocr: bool = True, enable_advanced_chunking: bool = True):
        self.extractors = {
            'pdf': PDFTextExtractor(enable_ocr=enable_ocr),
            'docx': DOCXTextExtractor(),
            'email': EmailTextExtractor()
        }
        
        self.chunker = SmartTextChunker() if enable_advanced_chunking else None
    
    async def process_document(self, 
                             file_path: str, 
                             file_type: str,
                             apply_advanced_chunking: bool = True) -> ExtractedDocument:
        """Process a document and extract structured text"""
        
        extractor = self.extractors.get(file_type.lower())
        if not extractor:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        if not extractor.validate_file(file_path):
            raise ValueError(f"Invalid file: {file_path}")
        
        # Extract text
        extracted_doc = await extractor.extract_text(file_path)
        
        # Apply advanced chunking if enabled
        if apply_advanced_chunking and self.chunker:
            extracted_doc = await self.chunker.rechunk_document(extracted_doc)
        
        logger.info(f"Document processing completed: {extracted_doc.document_id}")
        logger.info(f"Total chunks: {len(extracted_doc.chunks)}")
        logger.info(f"Total text length: {len(extracted_doc.full_text)} characters")
        
        return extracted_doc
    
    async def process_multiple_documents(self, 
                                       document_infos: List[Tuple[str, str]],  # [(file_path, file_type), ...]
                                       max_concurrent: int = 3) -> List[ExtractedDocument]:
        """Process multiple documents concurrently"""
        semaphore = asyncio.Semaphore(max_concurrent)
        
        async def bounded_process(file_path: str, file_type: str):
            async with semaphore:
                try:
                    return await self.process_document(file_path, file_type)
                except Exception as e:
                    logger.error(f"Failed to process {file_path}: {e}")
                    return None
        
        tasks = [bounded_process(fp, ft) for fp, ft in document_infos]
        results = await asyncio.gather(*tasks, return_exceptions=False)
        
        # Filter out None results
        successful_results = [r for r in results if r is not None]
        
        logger.info(f"Batch processing completed: {len(successful_results)}/{len(document_infos)} successful")
        return successful_results
    
    def save_extracted_document(self, extracted_doc: ExtractedDocument, output_path: str):
        """Save extracted document to JSON file"""
        # Convert to serializable format
        doc_dict = {
            "document_id": extracted_doc.document_id,
            "source_file": extracted_doc.source_file,
            "file_type": extracted_doc.file_type,
            "full_text": extracted_doc.full_text,
            "chunks": [
                {
                    "content": chunk.content,
                    "chunk_id": chunk.chunk_id,
                    "chunk_type": chunk.chunk_type,
                    "page_number": chunk.page_number,
                    "section_title": chunk.section_title,
                    "word_count": chunk.word_count,
                    "metadata": chunk.metadata
                }
                for chunk in extracted_doc.chunks
            ],
            "tables": extracted_doc.tables,
            "metadata": extracted_doc.metadata,
            "extraction_timestamp": extracted_doc.extraction_timestamp.isoformat(),
            "processing_stats": extracted_doc.processing_stats
        }
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(doc_dict, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Extracted document saved to: {output_path}")

# Usage example
async def main():
    """Example usage of the text extraction system"""
    
    # Initialize the processor
    processor = DocumentTextProcessor(enable_ocr=True, enable_advanced_chunking=True)
    
    # Example: Process a single document
    try:
        # Replace with actual file path from document ingestion
        file_path = "./document_cache/sample_policy.pdf"
        file_type = "pdf"
        
        # Process the document
        extracted_doc = await processor.process_document(file_path, file_type)
        
        print(f"Document ID: {extracted_doc.document_id}")
        print(f"Total text length: {len(extracted_doc.full_text)} characters")
        print(f"Number of chunks: {len(extracted_doc.chunks)}")
        print(f"Number of tables: {len(extracted_doc.tables)}")
        
        # Display first few chunks
        print("\nFirst 3 chunks:")
        for i, chunk in enumerate(extracted_doc.chunks[:3]):
            print(f"\nChunk {i+1} ({chunk.chunk_type}):")
            print(f"ID: {chunk.chunk_id}")
            if chunk.section_title:
                print(f"Section: {chunk.section_title}")
            print(f"Content preview: {chunk.content[:200]}...")
        
        # Save extracted document
        output_path = f"./extracted_{extracted_doc.document_id}.json"
        processor.save_extracted_document(extracted_doc, output_path)
        
        # Example: Process multiple documents
        document_batch = [
            ("./document_cache/policy1.pdf", "pdf"),
            ("./document_cache/contract1.docx", "docx"),
            ("./document_cache/email1.eml", "email")
        ]
        
        batch_results = await processor.process_multiple_documents(document_batch, max_concurrent=2)
        print(f"\nBatch processing completed: {len(batch_results)} documents processed")
        
    except Exception as e:
        print(f"Error processing document: {e}")

if __name__ == "__main__":
    asyncio.run(main())